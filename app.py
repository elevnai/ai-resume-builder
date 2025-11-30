from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import os
from openai import OpenAI
import PyPDF2
import docx
from io import BytesIO
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
MODEL = os.getenv('OPENAI_MODEL', 'gpt-4o-mini')

def extract_text_from_pdf(file_content):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_docx(file_content):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

def extract_text_from_file(file):
    """Extract text from uploaded file based on file type"""
    file_content = file.read()
    filename = file.filename.lower()

    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file_content)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(file_content)
    elif filename.endswith('.txt'):
        return file_content.decode('utf-8')
    else:
        return "Unsupported file format"

def tailor_resume(original_resume, job_description):
    """Use ChatGPT to tailor resume based on job description"""
    try:
        system_prompt = """You are an expert resume writer and career coach. Your task is to tailor a resume to match a specific job description while maintaining truthfulness and the candidate's actual experience.

Format your response as a structured resume with these EXACT sections in this order:

**PROFESSIONAL SUMMARY**
[3-4 sentences highlighting relevant experience and skills]

**EDUCATION & CERTIFICATIONS**
• [Certification/Degree] - [Date or Details]
• [Certification/Degree] - [Date or Details]

**PROFESSIONAL EXPERIENCE**

**[Job Title]** - [Company Name]  |  [Location]  |  [Dates]
• [Achievement/responsibility with metrics]
• [Achievement/responsibility with metrics]
• [Achievement/responsibility with metrics]

**[Job Title]** - [Company Name]  |  [Location]  |  [Dates]
• [Achievement/responsibility with metrics]
• [Achievement/responsibility with metrics]

**COMPUTER & SOFTWARE PROFICIENCY**
• [Tool 1]
• [Tool 2]
• [Tool 3]

Guidelines:
1. Analyze the job description to identify key skills, requirements, and keywords
2. Reorganize and rephrase the resume to highlight relevant experience
3. Use terminology and keywords from the job description
4. Keep all information truthful - only emphasize and reframe, never fabricate
5. Use bullet points (•) for all lists
6. Include metrics and specific achievements where possible
7. Optimize for ATS (Applicant Tracking Systems)
8. Keep the resume concise and impactful

Return ONLY the formatted resume text as shown above, ready to be copied or downloaded."""

        user_prompt = f"""Original Resume:
{original_resume}

Job Description:
{job_description}

Please tailor this resume to match the job description using the format specified."""

        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=2500
        )

        tailored_resume = response.choices[0].message.content
        return tailored_resume

    except Exception as e:
        raise Exception(f"Error tailoring resume: {str(e)}")

def create_formatted_docx(resume_text, name="Your Name"):
    """Create a professionally formatted DOCX resume"""
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Parse the resume text into sections
    lines = resume_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers (ALL CAPS sections)
        if line.isupper() and len(line) > 3:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.space_before = Pt(6)
            p.space_after = Pt(3)

        # Bold job titles or company lines
        elif '**' in line:
            p = doc.add_paragraph()
            # Remove ** markers and make bold
            clean_line = line.replace('**', '')
            run = p.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(10)
            p.space_before = Pt(3)

        # Bullet points
        elif line.startswith('•'):
            p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.runs[0]
            run.font.size = Pt(10)

        # Regular text
        else:
            p = doc.add_paragraph(line)
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.size = Pt(10)

    # Save to BytesIO
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)

    return docx_io

@app.route('/')
def index():
    """Serve the main page"""
    return render_template('index.html')

@app.route('/api/tailor-resume', methods=['POST'])
def tailor_resume_endpoint():
    """API endpoint to tailor resume"""
    try:
        # Check if file was uploaded
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file uploaded'}), 400

        resume_file = request.files['resume']
        job_description = request.form.get('job_description', '')

        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400

        # Extract text from resume
        resume_text = extract_text_from_file(resume_file)

        if resume_text.startswith('Error'):
            return jsonify({'error': resume_text}), 400

        # Tailor the resume
        tailored_resume = tailor_resume(resume_text, job_description)

        return jsonify({
            'success': True,
            'tailored_resume': tailored_resume
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-docx', methods=['POST'])
def download_docx():
    """Generate and download formatted DOCX resume"""
    try:
        data = request.get_json()
        resume_text = data.get('resume_text', '')

        if not resume_text:
            return jsonify({'error': 'No resume text provided'}), 400

        # Create formatted DOCX
        docx_file = create_formatted_docx(resume_text)

        return send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='tailored_resume.docx'
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'model': MODEL})

if __name__ == '__main__':
    port = int(os.getenv('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
